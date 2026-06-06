import os
import io
from flask import Blueprint, send_file, current_app
from app.models import Book, Page
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

export_bp = Blueprint('export', __name__)


def set_cell_background(cell, fill_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


@export_bp.route('/api/books/<int:book_id>/export/docx', methods=['GET'])
def export_docx(book_id):
    book = Book.query.get_or_404(book_id)
    pages = Page.query.filter_by(book_id=book_id).order_by(Page.page_number).all()

    doc = DocxDocument()

    # ── Document styles ──────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Cover / Title ────────────────────────────────────────────────────────
    title_para = doc.add_heading(book.title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    if book.author:
        author_para = doc.add_paragraph(f'Autor: {book.author}')
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_para = doc.add_paragraph(
        f'{book.source_language} → {book.target_language}  |  {len(pages)} páginas'
    )
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    upload_folder = current_app.config['UPLOAD_FOLDER']

    for page in pages:
        # ── Page heading ─────────────────────────────────────────────────────
        doc.add_heading(f'Página {page.page_number}', level=2)

        # ── Image ─────────────────────────────────────────────────────────────
        img_path = os.path.join(upload_folder, page.image_filename)
        if os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(6))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph(f'[Imagen no disponible: {page.image_filename}]')

        # ── Translation table ─────────────────────────────────────────────────
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = f'Original ({book.source_language})'
        hdr_cells[1].text = f'Traducción ({book.target_language})'
        for cell in hdr_cells:
            set_cell_background(cell, '1A1A2E')
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Content row
        content_cells = table.rows[1].cells
        content_cells[0].text = page.original_text or ''
        content_cells[1].text = page.translation or ''
        set_cell_background(content_cells[0], 'F0F4FF')
        set_cell_background(content_cells[1], 'FFF8F0')

        # Notes
        if page.notes and page.notes.strip():
            notes_para = doc.add_paragraph()
            run = notes_para.add_run(f'📝 Notas: {page.notes}')
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()  # spacer

        # Page break between pages (except last)
        if page != pages[-1]:
            doc.add_page_break()

    # ── Save to buffer ────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = book.slug
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f'{safe_title}_traduccion.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
